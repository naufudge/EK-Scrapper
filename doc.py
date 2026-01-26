from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from PIL import Image
from bs4 import NavigableString
from docx.text.paragraph import Paragraph
import io, os, httpx


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color and underline for hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # RTL for hyperlink
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

    # Faruma font for Complex Script (RTL text)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), 'Faruma')
    rPr.append(rFonts)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def create_initial_doc(doc_name: str, hijri_date: str, dhivehi_date: str):
    document = Document()
    core_properties = document.core_properties
    core_properties.author = 'National Archives of Maldives'
    core_properties.comments = ''
    
    # Styles
    styles = document.styles
    title_style = styles.add_style('headin', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Faruma'
    title_font.rtl = True
    # Set Complex Script font for RTL text
    rPr = title_style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Faruma')
    # Set Complex Script font size
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '32')  # Size in half-points (16pt = 32)
    rPr.append(szCs)

    author_style = styles.add_style('author', WD_STYLE_TYPE.PARAGRAPH)
    author_font = author_style.font
    author_font.name = 'Faruma'
    author_font.rtl = True
    # Set Complex Script font for RTL text
    rPr = author_style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Faruma')
    # Set Complex Script font size
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')  # Size in half-points (12pt = 24)
    rPr.append(szCs)

    date_style = styles.add_style('date_style', WD_STYLE_TYPE.PARAGRAPH)
    date_font = date_style.font
    date_font.name = 'Faruma'
    date_font.rtl = True
    # Set Complex Script font for RTL text
    rPr = date_style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Faruma')

    for date in [hijri_date, dhivehi_date]:
        head = document.add_paragraph(f"{date}")
        if date == hijri_date:
            head.paragraph_format.space_after = Pt(0)

        head.style = document.styles['date_style']
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(f'{doc_name}.docx')
    return document

def apply_rtl(paragraph: Paragraph):
    # 1. Set Paragraph Alignment (The easy way)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 2. Set Paragraph Properties (pPr) for RTL Direction
    pPr = paragraph._p.get_or_add_pPr()
    
    # This ensures the paragraph itself is flagged as RTL
    p_rtl = pPr.find(qn('w:rtl'))
    if p_rtl is None:
        p_rtl = OxmlElement('w:rtl')
        pPr.insert(0, p_rtl) # Insert at beginning for XML schema compliance

    # 3. Set Run Properties (rPr) for RTL Content
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        r_rtl = rPr.find(qn('w:rtl'))
        if r_rtl is None:
            r_rtl = OxmlElement('w:rtl')
            rPr.append(r_rtl)

def doc(filename, hijri_date, dhivehi_date, url: str, author: str, title: str, paras: list, image = None, update_url = True):
    if (os.path.exists(f"./{filename}.docx")):
        document = Document(f"{filename}.docx")
        try:
            # Styles
            styles = document.styles
            title_style = styles.add_style('headin', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Faruma'
            title_font.rtl = True
            # Set Complex Script font for RTL text
            rPr = title_style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'), 'Faruma')
            # Set Complex Script font size
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '32')  # Size in half-points (16pt = 32)
            rPr.append(szCs)

            author_style = styles.add_style('author', WD_STYLE_TYPE.PARAGRAPH)
            author_font = author_style.font
            author_font.name = 'Faruma'
            author_font.rtl = True
            # Set Complex Script font for RTL text
            rPr = author_style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'), 'Faruma')
            # Set Complex Script font size
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '24')  # Size in half-points (12pt = 24)
            rPr.append(szCs)

        except ValueError:
            pass
    else:
        document = create_initial_doc(filename, hijri_date, dhivehi_date)

    align_right = WD_ALIGN_PARAGRAPH.RIGHT
    align_center = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    main_head = document.add_heading(title.strip())
    # main_head.paragraph_format.space_after = Pt(2)
    main_head.style = document.styles['headin']
    main_head.alignment = align_right

    website_paragraph = document.add_paragraph()

    if (url.find("sun.mv") > 0):
        website = website_paragraph.add_run("ޚަބަރު ނެގީ: ސަން ވެބްސައިޓުން")
    elif (url.find("presidency.gov.mv") > 0):
        website = website_paragraph.add_run("ޚަބަރު ނެގީ: ރައީސް އޮފީސް ވެބްސައިޓުން")
    elif (url.find("mihaaru.com") > 0):
        website = website_paragraph.add_run("ޚަބަރު ނެގީ: މިހާރު ވެބްސައިޓުން")
    elif (url.find("avas.mv") > 0):
        website = website_paragraph.add_run("ޚަބަރު ނެގީ: އަވަސް ވެބްސައިޓުން")

    # Apply author style to website paragraph
    website_paragraph.style = document.styles['author']
    website_paragraph.paragraph_format.space_after = Pt(0)
    website_paragraph.alignment = align_right

    # Main body text style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Faruma'
    font.rtl = True
    # Set Complex Script font for RTL text
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Faruma')
    

    # Link
    link = document.add_paragraph("ލިންކް: \u200F\u202A " + url)
    link.paragraph_format.space_after = Pt(0) # No space after paragraph
    link.style = document.styles['author']
    link.alignment = align_right

    # Author
    if (author != ""):
        auth = document.add_paragraph("ލިޔުނީ:  " + author.strip())
        # auth.paragraph_format.space_after = Pt(1.5)
        auth.style = document.styles['author']
        auth.alignment = align_right
    else:
        auth = document.add_paragraph("ލިޔުނީ: ނެތް")
        # auth.paragraph_format.space_after = Pt(2)
        auth.style = document.styles['author']
        auth.alignment = align_right

    # Image
    if image != "" and image != None:
        timeout = httpx.Timeout(5, connect_timeout=None, read_timeout=None)
        with httpx.Client(http2=True, timeout=timeout) as client:
            imgLinkData = io.BytesIO(client.get(image).content)
            try:
                picture = document.add_picture(imgLinkData, width=Inches(2))
            except UnrecognizedImageError:
                # Picture not recognized by python-docx.
                img = Image.open(imgLinkData).convert("RGB")
                img_buffer = io.BytesIO()
                img.save(img_buffer, format="JPEG")
                picture = document.add_picture(img_buffer, width=Inches(2))

            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = align_center
            # last_paragraph.paragraph_format.space_after = Pt(2)

    # Paragraphs
    for each in paras:
        try:
            if each.find('span'):
                continue
            elif "relative" in each['class']:
                continue
        except KeyError:
            pass

        # Handle standalone <a> tags (not inside <p>)
        if each.name == 'a' and each.get('href'):
            main_body = document.add_paragraph()
            main_body.alignment = align_right
            link_text = each.get_text().strip()
            link_url = each['href']
            if link_text and link_url:
                add_hyperlink(main_body, link_url, link_text)
            apply_rtl(main_body)
            continue

        main_body = document.add_paragraph()
        main_body.alignment = align_right

        # Process paragraph contents preserving hyperlinks
        for child in each.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    main_body.add_run(text)
            elif child.name == 'a' and child.get('href'):
                link_text = child.get_text().strip()
                link_url = child['href']
                if link_text and link_url:
                    add_hyperlink(main_body, link_url, link_text)
            elif child.name:
                # Handle other nested elements - extract text
                text = child.get_text().strip()
                if text:
                    main_body.add_run(text)

        apply_rtl(main_body)

    document.add_page_break()

    document.save(f'{filename}.docx')

    if update_url:
        return f"'{url}' - Done"
    else:
        return f"'{title}' - Done"
